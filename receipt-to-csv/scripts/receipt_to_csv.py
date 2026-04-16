#!/usr/bin/env python3
"""Materialize receipt outputs from an agent-generated extraction payload."""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import io
import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore
    Inches = None  # type: ignore

try:
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover
    Image = None  # type: ignore

try:
    import requests  # type: ignore
except Exception:  # pragma: no cover
    requests = None  # type: ignore

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover
    fitz = None  # type: ignore


SUPPORTED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif", ".pdf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif"}
CSV_HEADERS = [
    "No",
    "Original Filename",
    "New Name",
    "Receipt Date",
    "Merchant",
    "Description",
    "Rcpt Ccy",
    "Rcpt Amt",
    "Reimb Ccy",
    "Reimb Amt",
    "Remark",
    "Category",
    "Status",
]
STATUS_VALUES = {"OK", "NEEDS_REVIEW", "FAILED"}
FX_API_BASE = "https://api.frankfurter.dev/v1"


@dataclass
class ReceiptRow:
    no: int
    original_filename: str
    new_name: str
    receipt_date: str
    merchant: str
    description: str
    receipt_currency: str
    receipt_amount: Optional[float]
    reimb_currency: str
    reimb_amount: Optional[float]
    remark: str
    category: str
    status: str
    copied_path: Path


def sanitize_filename(value: str, fallback: str = "UNKNOWN") -> str:
    cleaned = value.strip().replace("\n", " ")
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", cleaned)
    cleaned = re.sub(r"\s+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    if not cleaned:
        cleaned = fallback
    return cleaned[:120]


def normalize_date(raw: Any) -> str:
    if raw is None:
        return ""
    text = str(raw).strip()
    if not text:
        return ""
    text = text.replace(".", "-").replace("/", "-")
    match = re.fullmatch(r"(\d{4})-(\d{1,2})-(\d{1,2})", text)
    if not match:
        return text
    year, month, day = match.groups()
    return f"{year}-{int(month):02d}-{int(day):02d}"


def normalize_currency(raw: Any) -> str:
    if raw is None:
        return ""
    text = str(raw).strip().upper()
    return text if re.fullmatch(r"[A-Z]{3}", text) else ""


def to_float(raw: Any) -> Optional[float]:
    if raw is None or raw == "":
        return None
    if isinstance(raw, (int, float)):
        return float(raw)
    text = str(raw).strip().replace(",", "")
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def get_first(entry: dict[str, Any], keys: list[str], default: Any = "") -> Any:
    for key in keys:
        if key in entry and entry[key] is not None:
            return entry[key]
    return default


def append_remark(remark: str, note: str) -> str:
    note = note.strip()
    if not note:
        return remark
    if not remark:
        return note
    return f"{remark}; {note}"


def ensure_output_dirs(source_folder: Path, output_subfolder: Optional[str]) -> tuple[Path, Path]:
    if output_subfolder:
        folder_name = sanitize_filename(output_subfolder, "receipt_to_csv_output")
    else:
        folder_name = f"receipt_to_csv_output_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}"

    output_dir = source_folder / folder_name
    if output_dir.exists():
        raise FileExistsError(f"Output folder already exists: {output_dir}")
    attachments_dir = output_dir / "attachments"
    attachments_dir.mkdir(parents=True, exist_ok=False)
    return output_dir, attachments_dir


def load_payload(payload_file: Optional[str], payload_json: Optional[str]) -> list[dict[str, Any]]:
    payload: Any
    if payload_file:
        payload_text = Path(payload_file).read_text(encoding="utf-8")
        payload = json.loads(payload_text)
    elif payload_json:
        payload = json.loads(payload_json)
    else:
        raise ValueError("Either --payload-file or --payload-json is required.")

    if isinstance(payload, list):
        rows = payload
    elif isinstance(payload, dict) and isinstance(payload.get("rows"), list):
        rows = payload["rows"]
    else:
        raise ValueError("Payload must be a JSON array or an object with a 'rows' array.")

    normalized: list[dict[str, Any]] = []
    for idx, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            raise ValueError(f"Payload row {idx} is not an object.")
        normalized.append(row)
    return normalized


def resolve_source_file(
    source_folder: Path,
    entry: dict[str, Any],
    fallback_candidates: dict[str, list[Path]],
) -> Optional[Path]:
    source_path_raw = get_first(entry, ["source_path", "sourcePath"], "")
    if source_path_raw:
        candidate = Path(str(source_path_raw))
        if not candidate.is_absolute():
            candidate = (source_folder / candidate).resolve()
        if candidate.exists() and candidate.is_file():
            return candidate
        return None

    original_name = str(get_first(entry, ["original_filename", "originalFilename"], "")).strip()
    if original_name and original_name in fallback_candidates:
        candidates = fallback_candidates[original_name]
        if len(candidates) == 1:
            return candidates[0]

    return None


def build_new_name(index: int, source: Path, entry: dict[str, Any], used_names: set[str]) -> str:
    explicit = str(get_first(entry, ["new_name", "newName"], "")).strip()
    ext = source.suffix.lower()
    if explicit:
        explicit = sanitize_filename(Path(explicit).stem, fallback=f"{index:03d}_RECEIPT")
        base_name = explicit
    else:
        date_part = normalize_date(get_first(entry, ["receipt_date", "receiptDate"], "")) or "0000-00-00"
        merchant_part = sanitize_filename(str(get_first(entry, ["merchant"], "UNKNOWN_MERCHANT")), "UNKNOWN_MERCHANT")
        amount = to_float(get_first(entry, ["receipt_amount", "receiptAmount"], None))
        amount_part = f"{amount:.2f}" if amount is not None else "0.00"
        ccy_part = normalize_currency(get_first(entry, ["receipt_currency", "receiptCurrency"], "UNK")) or "UNK"
        base_name = sanitize_filename(
            f"{index:03d}_{date_part}_{merchant_part}_{amount_part}_{ccy_part}",
            fallback=f"{index:03d}_RECEIPT",
        )

    name = f"{base_name}{ext}"
    counter = 1
    while name in used_names:
        name = f"{base_name}_{counter:02d}{ext}"
        counter += 1
    used_names.add(name)
    return name


def format_amount(value: Optional[float]) -> str:
    return "" if value is None else f"{value:.2f}"


def fetch_fx_rate(
    from_currency: str,
    to_currency: str,
    receipt_date: str,
    fx_cache: dict[tuple[str, str, str], float],
) -> tuple[Optional[float], Optional[str]]:
    if from_currency == to_currency:
        return 1.0, None
    if requests is None:
        return None, "requests package is missing for FX lookup"

    normalized_date = normalize_date(receipt_date)
    date_segment = normalized_date if re.fullmatch(r"\d{4}-\d{2}-\d{2}", normalized_date) else "latest"

    cache_key = (from_currency, to_currency, date_segment)
    if cache_key in fx_cache:
        return fx_cache[cache_key], None

    def _query(segment: str) -> tuple[Optional[float], Optional[str]]:
        url = f"{FX_API_BASE}/{segment}"
        try:
            response = requests.get(url, params={"from": from_currency, "to": to_currency}, timeout=20)
        except Exception as error:
            return None, f"FX request failed: {error}"
        if not response.ok:
            return None, f"FX API HTTP {response.status_code}"
        try:
            payload = response.json()
        except Exception:
            return None, "FX API JSON parse failed"
        rates = payload.get("rates", {})
        rate = rates.get(to_currency)
        if not isinstance(rate, (int, float)) or rate <= 0:
            return None, f"No FX rate for {from_currency}->{to_currency}"
        return float(rate), None

    rate, error = _query(date_segment)
    if rate is None and date_segment != "latest":
        rate, error = _query("latest")

    if rate is not None:
        fx_cache[cache_key] = rate
    return rate, error


def compute_reimbursement_amount(
    receipt_amount: Optional[float],
    receipt_currency: str,
    reimb_currency: str,
    receipt_date: str,
    rate_hint: Optional[float],
    fx_cache: dict[tuple[str, str, str], float],
) -> tuple[Optional[float], Optional[str], Optional[float]]:
    if receipt_amount is None:
        return None, None, None
    if not receipt_currency:
        return None, "missing receipt currency", None
    if not reimb_currency:
        return None, "missing reimbursement currency", None

    if receipt_currency == reimb_currency:
        return receipt_amount, None, 1.0

    if rate_hint is not None and rate_hint > 0:
        return receipt_amount * rate_hint, None, rate_hint

    rate, error = fetch_fx_rate(receipt_currency, reimb_currency, receipt_date, fx_cache)
    if rate is None:
        return None, error or "FX lookup failed", None
    return receipt_amount * rate, None, rate


def build_status(
    receipt_date: str,
    merchant: str,
    receipt_amount: Optional[float],
    conversion_ok: bool,
    explicit: str,
) -> str:
    explicit_normalized = explicit.strip().upper()
    if explicit_normalized in STATUS_VALUES:
        return explicit_normalized
    if not receipt_date or not merchant or receipt_amount is None or not conversion_ok:
        return "NEEDS_REVIEW"
    return "OK"


def write_csv_report(rows: list[ReceiptRow], csv_path: Path) -> None:
    with csv_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=CSV_HEADERS)
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "No": row.no,
                    "Original Filename": row.original_filename,
                    "New Name": row.new_name,
                    "Receipt Date": row.receipt_date,
                    "Merchant": row.merchant,
                    "Description": row.description,
                    "Rcpt Ccy": row.receipt_currency,
                    "Rcpt Amt": format_amount(row.receipt_amount),
                    "Reimb Ccy": row.reimb_currency,
                    "Reimb Amt": format_amount(row.reimb_amount),
                    "Remark": row.remark,
                    "Category": row.category,
                    "Status": row.status,
                }
            )


def pdf_page_count(pdf_path: Path) -> Optional[int]:
    if shutil.which("pdfinfo") is None:
        return None
    try:
        result = subprocess.run(
            ["pdfinfo", str(pdf_path)],
            capture_output=True,
            text=True,
            check=False,
        )
    except Exception:
        return None
    if result.returncode != 0:
        return None
    match = re.search(r"^Pages:\s+(\d+)\s*$", result.stdout, flags=re.MULTILINE)
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def render_pdf_pages_via_pymupdf(
    pdf_path: Path,
    max_pdf_pages: int,
) -> tuple[list[bytes], bool, Optional[str]]:
    if fitz is None:
        return [], False, "PyMuPDF not installed"
    pages: list[bytes] = []
    truncated = False
    try:
        with fitz.open(str(pdf_path)) as pdf_doc:
            total_pages = len(pdf_doc)
            limit = total_pages if max_pdf_pages <= 0 else min(total_pages, max_pdf_pages)
            truncated = total_pages > limit
            for page_index in range(limit):
                page = pdf_doc[page_index]
                pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                pages.append(pix.tobytes("jpg"))
    except Exception as error:
        return [], False, str(error)
    return pages, truncated, None


def render_pdf_pages_via_pdftoppm(
    pdf_path: Path,
    max_pdf_pages: int,
) -> tuple[list[bytes], bool, Optional[str]]:
    if shutil.which("pdftoppm") is None:
        return [], False, "pdftoppm not installed"

    with tempfile.TemporaryDirectory(prefix="receipt_pdf_") as temp_dir:
        temp_root = Path(temp_dir)
        prefix = temp_root / "page"
        command = ["pdftoppm", "-jpeg", "-jpegopt", "quality=85"]
        if max_pdf_pages > 0:
            command.extend(["-f", "1", "-l", str(max_pdf_pages)])
        command.extend([str(pdf_path), str(prefix)])

        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                check=False,
            )
        except Exception as error:
            return [], False, str(error)

        if result.returncode != 0:
            error_text = result.stderr.strip() or result.stdout.strip() or f"exit code {result.returncode}"
            return [], False, f"pdftoppm failed: {error_text}"

        image_paths = sorted(
            temp_root.glob("page-*.jpg"),
            key=lambda path: int(path.stem.split("-")[-1]) if path.stem.split("-")[-1].isdigit() else 0,
        )
        if not image_paths:
            return [], False, "pdftoppm produced no images"

        pages = [path.read_bytes() for path in image_paths]
        total_pages = pdf_page_count(pdf_path)
        truncated = bool(max_pdf_pages > 0 and total_pages is not None and total_pages > max_pdf_pages)
        return pages, truncated, None


def render_pdf_pages_via_qlmanage(
    pdf_path: Path,
) -> tuple[list[bytes], bool, Optional[str]]:
    if shutil.which("qlmanage") is None:
        return [], False, "qlmanage not installed"

    with tempfile.TemporaryDirectory(prefix="receipt_pdf_ql_") as temp_dir:
        try:
            result = subprocess.run(
                ["qlmanage", "-t", "-s", "2200", "-o", temp_dir, str(pdf_path)],
                capture_output=True,
                text=True,
                check=False,
            )
        except Exception as error:
            return [], False, str(error)

        if result.returncode != 0:
            error_text = result.stderr.strip() or result.stdout.strip() or f"exit code {result.returncode}"
            return [], False, f"qlmanage failed: {error_text}"

        image_paths = sorted(Path(temp_dir).glob("*.png"))
        if not image_paths:
            return [], False, "qlmanage produced no preview image"

        # qlmanage generates a thumbnail, usually first page only.
        return [image_paths[0].read_bytes()], True, None


def render_pdf_pages_as_images(
    pdf_path: Path,
    max_pdf_pages: int,
) -> tuple[list[bytes], bool, Optional[str]]:
    attempts: list[str] = []

    pages, truncated, error = render_pdf_pages_via_pymupdf(pdf_path, max_pdf_pages)
    if pages:
        return pages, truncated, None
    if error:
        attempts.append(f"pymupdf: {error}")

    pages, truncated, error = render_pdf_pages_via_pdftoppm(pdf_path, max_pdf_pages)
    if pages:
        return pages, truncated, None
    if error:
        attempts.append(f"pdftoppm: {error}")

    pages, truncated, error = render_pdf_pages_via_qlmanage(pdf_path)
    if pages:
        return pages, truncated, None
    if error:
        attempts.append(f"qlmanage: {error}")

    combined = "; ".join(attempts) if attempts else "no PDF rendering backend available"
    return [], False, combined


def write_docx_report(rows: list[ReceiptRow], docx_path: Path, max_pdf_pages: int) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

    doc = Document()
    doc.add_heading("ReceiptToCSV Report", level=0)
    doc.add_paragraph(f"Generated on {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Receipt Date"
    header[2].text = "Filename"
    header[3].text = "Merchant / Description"
    header[4].text = "Reimb Ccy"
    header[5].text = "Reimb Amt"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.no)
        cells[1].text = row.receipt_date
        cells[2].text = row.new_name
        cells[3].text = f"{row.merchant} - {row.description}".strip(" -")
        cells[4].text = row.reimb_currency
        cells[5].text = format_amount(row.reimb_amount)

    doc.add_page_break()
    doc.add_heading("Attachments", level=1)
    for row in rows:
        doc.add_paragraph(f"{row.no}. {row.new_name}", style="List Number")
        ext = row.copied_path.suffix.lower()

        if ext in IMAGE_EXTENSIONS and row.copied_path.exists() and Inches is not None:
            try:
                if ext in {".webp", ".heic", ".heif"} and Image is not None:
                    with Image.open(row.copied_path) as img:
                        rgb_img = img.convert("RGB")
                        image_bytes = io.BytesIO()
                        rgb_img.save(image_bytes, format="JPEG")
                        image_bytes.seek(0)
                        doc.add_picture(image_bytes, width=Inches(5.8))
                else:
                    doc.add_picture(str(row.copied_path), width=Inches(5.8))
            except Exception:
                doc.add_paragraph("[Image preview unavailable]")
            continue

        if ext == ".pdf" and Inches is not None:
            pdf_images, truncated, pdf_error = render_pdf_pages_as_images(row.copied_path, max_pdf_pages)
            if pdf_images:
                for page_index, image_data in enumerate(pdf_images, start=1):
                    if len(pdf_images) > 1:
                        doc.add_paragraph(f"Page {page_index} of {len(pdf_images)}")
                    doc.add_picture(io.BytesIO(image_data), width=Inches(5.8))
                if truncated and max_pdf_pages > 0:
                    doc.add_paragraph(f"[Only first {max_pdf_pages} pages were embedded]")
            else:
                doc.add_paragraph(f"[Attachment file: {row.new_name}]")
                if pdf_error:
                    doc.add_paragraph(f"[PDF preview unavailable: {pdf_error}]")
            continue

        doc.add_paragraph(f"[Attachment file: {row.new_name}]")

    doc.save(docx_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create report.csv, report.docx, and attachments/ from a structured payload."
    )
    parser.add_argument("folder", nargs="?", help="Folder path that contains source receipts")
    parser.add_argument(
        "--payload-file",
        default=None,
        help="Path to JSON payload file (array or object with rows array)",
    )
    parser.add_argument(
        "--payload-json",
        default=None,
        help="Inline JSON payload string (array or object with rows array)",
    )
    parser.add_argument(
        "--home-currency",
        default="SGD",
        help="Home reimbursement currency for conversion (default: SGD)",
    )
    parser.add_argument(
        "--reimb-currency",
        default=None,
        help="Deprecated alias of --home-currency (kept for compatibility)",
    )
    parser.add_argument(
        "--output-subfolder",
        default=None,
        help="Custom output subfolder name (default: receipt_to_csv_output_<timestamp>)",
    )
    parser.add_argument(
        "--max-pdf-pages",
        type=int,
        default=8,
        help="Maximum PDF pages to embed per file in DOCX (0 means all, default: 8)",
    )
    return parser.parse_args()


def prompt_for_folder() -> str:
    return input("Which folder path should I scan for receipts (images and PDFs)? ").strip()


def build_fallback_candidates(source_folder: Path) -> dict[str, list[Path]]:
    candidates: dict[str, list[Path]] = {}
    for file in source_folder.rglob("*"):
        if not file.is_file():
            continue
        if file.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        candidates.setdefault(file.name, []).append(file.resolve())
    return candidates


def main() -> int:
    args = parse_args()

    if Document is None:
        print("Missing dependency: python-docx. Install with `pip install python-docx`.")
        return 2

    if not args.payload_file and not args.payload_json:
        print("Missing payload. Use --payload-file or --payload-json.")
        return 2

    folder_text = args.folder or prompt_for_folder()
    if not folder_text:
        print("No folder path was provided.")
        return 2

    source_folder = Path(folder_text).expanduser().resolve()
    if not source_folder.exists() or not source_folder.is_dir():
        print(f"Invalid folder path: {source_folder}")
        return 2

    try:
        payload_rows = load_payload(args.payload_file, args.payload_json)
    except Exception as error:
        print(f"Invalid payload: {error}")
        return 2

    if not payload_rows:
        print("Payload has no rows.")
        return 2

    try:
        output_dir, attachments_dir = ensure_output_dirs(source_folder, args.output_subfolder)
    except FileExistsError as error:
        print(str(error))
        return 2

    home_currency = normalize_currency(args.home_currency or args.reimb_currency or "SGD") or "SGD"
    used_names: set[str] = set()
    rows: list[ReceiptRow] = []
    skipped = 0
    fallback_candidates = build_fallback_candidates(source_folder)
    fx_cache: dict[tuple[str, str, str], float] = {}

    for idx, entry in enumerate(payload_rows, start=1):
        source_file = resolve_source_file(source_folder, entry, fallback_candidates)
        if source_file is None:
            skipped += 1
            print(f"[{idx}/{len(payload_rows)}] Skipped: could not resolve source file")
            continue
        if source_file.suffix.lower() not in SUPPORTED_EXTENSIONS:
            skipped += 1
            print(f"[{idx}/{len(payload_rows)}] Skipped unsupported file type: {source_file.name}")
            continue

        new_name = build_new_name(idx, source_file, entry, used_names)
        copied_path = attachments_dir / new_name
        shutil.copy2(source_file, copied_path)

        receipt_date = normalize_date(get_first(entry, ["receipt_date", "receiptDate"], ""))
        merchant = str(get_first(entry, ["merchant"], "")).strip()
        description = str(get_first(entry, ["description"], "")).strip()
        rcpt_currency = normalize_currency(get_first(entry, ["receipt_currency", "receiptCurrency"], ""))
        rcpt_amount = to_float(get_first(entry, ["receipt_amount", "receiptAmount"], None))

        row_reimb_currency = normalize_currency(get_first(entry, ["reimb_currency", "reimbCurrency"], ""))
        reimb_currency = row_reimb_currency or home_currency
        reimb_amount = to_float(get_first(entry, ["reimb_amount", "reimbAmount"], None))
        rate_hint = to_float(get_first(entry, ["fx_rate_to_home", "fxRateToHome", "fx_rate", "fxRate"], None))
        remark = str(get_first(entry, ["remark"], "")).strip()

        fx_rate_used: Optional[float] = None
        fx_error: Optional[str] = None
        if reimb_amount is None:
            reimb_amount, fx_error, fx_rate_used = compute_reimbursement_amount(
                receipt_amount=rcpt_amount,
                receipt_currency=rcpt_currency,
                reimb_currency=reimb_currency,
                receipt_date=receipt_date,
                rate_hint=rate_hint,
                fx_cache=fx_cache,
            )
            if fx_error:
                remark = append_remark(remark, f"FX conversion failed: {fx_error}")
            elif fx_rate_used is not None and rcpt_currency and rcpt_currency != reimb_currency:
                remark = append_remark(remark, f"FX rate {rcpt_currency}->{reimb_currency}: {fx_rate_used:.6f}")

        category = str(get_first(entry, ["category"], "RECEIPT")).strip() or "RECEIPT"
        conversion_ok = not (rcpt_amount is not None and reimb_amount is None)
        status = build_status(
            receipt_date=receipt_date,
            merchant=merchant,
            receipt_amount=rcpt_amount,
            conversion_ok=conversion_ok,
            explicit=str(get_first(entry, ["status"], "")),
        )

        row = ReceiptRow(
            no=len(rows) + 1,
            original_filename=source_file.name,
            new_name=new_name,
            receipt_date=receipt_date,
            merchant=merchant,
            description=description,
            receipt_currency=rcpt_currency,
            receipt_amount=rcpt_amount,
            reimb_currency=reimb_currency,
            reimb_amount=reimb_amount,
            remark=remark,
            category=category,
            status=status,
            copied_path=copied_path,
        )
        rows.append(row)
        print(f"[{idx}/{len(payload_rows)}] Added: {source_file.name} -> {new_name}")

    if not rows:
        print("No rows were generated. Nothing to export.")
        return 2

    csv_path = output_dir / "report.csv"
    docx_path = output_dir / "report.docx"
    write_csv_report(rows, csv_path)
    write_docx_report(rows, docx_path, max_pdf_pages=args.max_pdf_pages)

    needs_review = sum(1 for row in rows if row.status != "OK")
    print("")
    print("Completed receipt export bundle:")
    print(f"- Home currency: {home_currency}")
    print(f"- Output folder: {output_dir}")
    print(f"- CSV: {csv_path}")
    print(f"- DOCX: {docx_path}")
    print(f"- Attachments folder: {attachments_dir}")
    print(f"- Processed rows: {len(rows)}")
    print(f"- Skipped rows: {skipped}")
    print(f"- NEEDS_REVIEW rows: {needs_review}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
